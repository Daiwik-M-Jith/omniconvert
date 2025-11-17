from io import BytesIO

from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfReader

from .base import register_converter


def _text_from_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def _text_to_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    for line in text.splitlines() or [""]:
        pdf.multi_cell(0, 8, line if line else " ")
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _text_to_png_bytes(text: str) -> bytes:
    font = ImageFont.load_default()
    lines = text.splitlines() or [""]
    width = int(max(font.getlength(line) for line in lines) + 24)
    height = (font.size + 6) * len(lines) + 20
    image = Image.new("RGB", (max(width, 200), max(height, 80)), color="white")
    draw = ImageDraw.Draw(image)
    y = 10
    for line in lines:
        draw.text((12, y), line, fill="black", font=font)
        y += font.size + 6
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


@register_converter("docx", "txt", note="Extracts plain text from DOCX")
def docx_to_txt(content: bytes, target: str = "txt"):
    text = _text_from_docx(content)
    return text.encode("utf-8"), "text/plain"


@register_converter("docx", "pdf", note="Text-only render via FPDF")
def docx_to_pdf(content: bytes, target: str = "pdf"):
    text = _text_from_docx(content)
    return _text_to_pdf_bytes(text), "application/pdf"


@register_converter("docx", "png", note="Text rendered into PNG (no layout)")
def docx_to_png(content: bytes, target: str = "png"):
    text = _text_from_docx(content)
    return _text_to_png_bytes(text), "image/png"


@register_converter("pdf", "txt", note="Extracts text via PyPDF2")
def pdf_to_txt(content: bytes, target: str = "txt"):
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages)
    return text.encode("utf-8"), "text/plain"


@register_converter("pdf", "docx", note="Creates DOCX with extracted text")
def pdf_to_docx(content: bytes, target: str = "docx"):
    reader = PdfReader(BytesIO(content))
    doc = Document()
    for page in reader.pages:
        doc.add_paragraph(page.extract_text() or "")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@register_converter("txt", "pdf", note="Wraps plaintext into PDF")
def txt_to_pdf(content: bytes, target: str = "pdf"):
    text = content.decode("utf-8", errors="ignore")
    return _text_to_pdf_bytes(text), "application/pdf"


@register_converter("txt", "png", note="Renders plaintext into PNG")
def txt_to_png(content: bytes, target: str = "png"):
    text = content.decode("utf-8", errors="ignore")
    return _text_to_png_bytes(text), "image/png"
